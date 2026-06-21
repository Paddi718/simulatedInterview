import uuid
import os
from datetime import datetime
from pathlib import Path
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from app.models.interview import Interview
from app.models.interview_question import InterviewQuestion
from app.models.interview_document import InterviewDocument
from app.models.resume import Resume
from app.models.job_description import JobDescription

QUESTION_TYPE_LABELS = {
    "introduction": "自我介绍",
    "behavioral": "行为面试",
    "technical": "专业技能",
    "situational": "情景题",
    "career": "职业规划",
}


def _build_markdown(interview: Interview, questions: list[InterviewQuestion]) -> str:
    """生成 Markdown 格式报告（使用 Jinja2 模板）"""
    from jinja2 import Environment, FileSystemLoader

    env = Environment(loader=FileSystemLoader(Path(__file__).parent.parent / "templates"))
    template = env.get_template("report.md")

    context = _build_context(interview, questions)
    return template.render(**context)


def _build_html(interview: Interview, questions: list[InterviewQuestion]) -> str:
    """生成 HTML 格式报告"""
    import markdown
    md = _build_markdown(interview, questions)
    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>面试报告</title>
    <style>
        body {{ font-family: 'Microsoft YaHei', sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f5f5f5; }}
        h1 {{ color: #1a1a1a; }}
        h2 {{ color: #333; border-bottom: 2px solid #eee; padding-bottom: 5px; }}
        h3 {{ color: #555; }}
    </style>
</head>
<body>
{markdown.markdown(md, extensions=['extra'])}
</body>
</html>"""
    return html


def _find_chinese_font() -> str | None:
    """查找系统中可用的中文字体"""
    import platform
    candidates = []
    if platform.system() == 'Windows':
        candidates = [
            'C:/Windows/Fonts/msyh.ttc',
            'C:/Windows/Fonts/simsun.ttc',
            'C:/Windows/Fonts/simhei.ttf',
        ]
    else:
        candidates = [
            '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',
            '/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf',
            '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
        ]
    for path in candidates:
        if os.path.exists(path):
            return path
    return None


def _build_pdf(interview: Interview, questions: list[InterviewQuestion]) -> bytes:
    """生成 PDF 格式报告
    1. 优先用 weasyprint（Linux/Docker）
    2. 其次用 fpdf2（跨平台，纯 Python）
    3. 都不行则返回 HTML
    """
    import shutil

    # 方式 1：weasyprint（需要 fontconfig）
    if shutil.which('fc-list'):
        try:
            from weasyprint import HTML as WHTML
            import concurrent.futures

            html = _build_html(interview, questions)

            def _render():
                return WHTML(string=html).write_pdf()

            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                return executor.submit(_render).result(timeout=15)
        except Exception:
            pass

    # 方式 2：fpdf2（纯 Python，支持中文）
    font_path = _find_chinese_font()
    if font_path:
        try:
            return _build_pdf_fpdf2(interview, questions, font_path)
        except Exception as e:
            import traceback
            traceback.print_exc()

    # 方式 3：降级为 HTML
    return _build_html(interview, questions).encode('utf-8')


def _build_pdf_fpdf2(interview: Interview, questions: list[InterviewQuestion], font_path: str) -> bytes:
    """使用 fpdf2 生成专业排版的 PDF 报告"""
    from fpdf import FPDF

    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_left_margin(20)
    pdf.set_right_margin(20)
    pdf.add_page()

    # Register fonts
    pdf.add_font('cjk', '', font_path)
    pdf.add_font('cjk', 'B', font_path)

    # Color palette
    PRIMARY = (25, 118, 210)     # Blue
    DARK = (33, 33, 33)
    GRAY = (117, 117, 117)
    LIGHT_BG = (245, 247, 250)
    ACCENT_BG = (232, 245, 253)
    GREEN_BG = (232, 253, 245)
    ORANGE_BG = (255, 248, 225)
    WHITE = (255, 255, 255)
    BORDER = (224, 224, 224)

    scores = interview.dimension_scores or {}
    page_w = pdf.w - pdf.l_margin - pdf.r_margin  # usable width

    # ---- Helper functions ----
    def section_title(text: str):
        pdf.ln(6)
        pdf.set_fill_color(*PRIMARY)
        pdf.set_text_color(*WHITE)
        pdf.set_font('cjk', 'B', 13)
        pdf.cell(page_w, 9, f'  {text}', fill=True, new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(*DARK)
        pdf.ln(5)

    def key_value(key: str, value: str, w_key: float = 42):
        pdf.set_font('cjk', '', 10)
        pdf.set_text_color(*GRAY)
        pdf.cell(w_key, 7, key)
        pdf.set_text_color(*DARK)
        pdf.set_font('cjk', 'B', 10)
        pdf.cell(page_w - w_key, 7, str(value), new_x="LMARGIN", new_y="NEXT")

    def body_text(text: str, size: float = 9.5):
        if not text:
            return
        pdf.set_font('cjk', '', size)
        pdf.set_text_color(*DARK)
        pdf.multi_cell(page_w, 5.5, text, align='L')
        pdf.ln(1.5)

    def label_text(label: str, size: float = 9):
        pdf.set_font('cjk', 'B', size)
        pdf.set_text_color(*DARK)
        pdf.cell(page_w, 6, label, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)

    def divider():
        pdf.ln(2)
        pdf.set_draw_color(*BORDER)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(3)

    # ====== Cover / Title ======
    pdf.ln(8)
    pdf.set_font('cjk', 'B', 26)
    pdf.set_text_color(*PRIMARY)
    pdf.cell(page_w, 14, '模拟面试报告', align='C', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)
    pdf.set_font('cjk', '', 11)
    pdf.set_text_color(*GRAY)
    pdf.cell(page_w, 7, f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}", align='C', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(10)

    # ====== Overview Card ======
    pdf.set_fill_color(*LIGHT_BG)
    pdf.set_draw_color(*BORDER)
    y0 = pdf.get_y()
    pdf.rect(pdf.l_margin, y0, page_w, 38, style='DF')
    pdf.set_xy(pdf.l_margin + 4, y0 + 4)

    # Score circle (fake — just big text)
    pdf.set_font('cjk', 'B', 42)
    pdf.set_text_color(*PRIMARY)
    pdf.cell(36, 18, str(interview.total_score or '-'), align='C')
    pdf.set_font('cjk', '', 9)
    pdf.set_text_color(*GRAY)
    pdf.cell(0, 18, '  分 / 100', new_x="LMARGIN", new_y="NEXT")

    pdf.set_x(pdf.l_margin + 4)
    pdf.set_font('cjk', '', 9.5)
    pdf.set_text_color(*DARK)
    diff_label = {'easy': '初级', 'mid': '中级', 'hard': '高级'}.get(interview.difficulty, interview.difficulty)
    pdf.cell(page_w - 8, 6, f"难度级别：{diff_label}　|　面试时间：{interview.started_at.strftime('%Y-%m-%d %H:%M') if interview.started_at else 'N/A'}", new_x="LMARGIN", new_y="NEXT")
    pdf.set_y(y0 + 38 + 6)

    # ====== Dimension Scores Table ======
    section_title('各维度评分')
    dims = [
        ('内容完整性', scores.get('content_completeness', '-')),
        ('专业度', scores.get('professionalism', '-')),
        ('表达能力', scores.get('expression', '-')),
        ('STAR 法则', scores.get('star_method', '-')),
    ]
    col_w = page_w / 4
    # Header row
    pdf.set_fill_color(*PRIMARY)
    pdf.set_text_color(*WHITE)
    pdf.set_font('cjk', 'B', 10)
    for label, _ in dims:
        pdf.cell(col_w, 10, label, border=0, fill=True, align='C')
    pdf.ln()
    # Value row
    pdf.set_fill_color(*LIGHT_BG)
    pdf.set_text_color(*DARK)
    pdf.set_font('cjk', 'B', 16)
    for _, val in dims:
        pdf.cell(col_w, 14, str(val), border=0, fill=True, align='C')
    pdf.ln(8)

    # ====== AI Overview ======
    if interview.ai_overview:
        section_title('综合评价')
        body_text(interview.ai_overview)

    # ====== Question Details ======
    section_title('逐题详情')
    sorted_qs = sorted(questions, key=lambda x: x.order_index)
    for idx, q in enumerate(sorted_qs):
        sd = q.score_detail or {}
        # Check if we need a page break (estimate: ~60mm per question)
        if pdf.get_y() > pdf.h - 75:
            pdf.add_page()

        # Question header bar
        q_num = f"第 {q.order_index} 题"
        type_label = QUESTION_TYPE_LABELS.get(q.question_type, q.question_type)
        pdf.set_fill_color(*PRIMARY)
        pdf.set_text_color(*WHITE)
        pdf.set_font('cjk', 'B', 10)
        pdf.cell(page_w, 8, f'  {q_num}  ·  {type_label}', fill=True, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

        # Question text
        pdf.set_font('cjk', 'B', 10)
        pdf.set_text_color(*DARK)
        pdf.multi_cell(page_w, 5.8, q.question_text)
        pdf.ln(1)

        # Score (if available)
        if q.ai_score is not None:
            pdf.set_fill_color(*LIGHT_BG)
            pdf.set_draw_color(*BORDER)
            y1 = pdf.get_y()
            pdf.rect(pdf.l_margin, y1, page_w, 10, style='DF')
            pdf.set_xy(pdf.l_margin + 3, y1 + 1.5)
            pdf.set_font('cjk', '', 9)
            pdf.set_text_color(*DARK)
            dim_str = '  |  '.join(
                f"{'内容' if k=='content_completeness' else '专业' if k=='professionalism' else '表达' if k=='expression' else 'STAR'}：{v}"
                for k, v in sd.items()
            )
            pdf.cell(page_w - 6, 7, f"得分：{q.ai_score} 分　　　{dim_str}")
            pdf.set_y(y1 + 10 + 3)

        # Answer
        label_text('▎你的回答', 9)
        body_text(q.user_answer_transcript or '（未作答）', 9)

        # Evaluation
        if q.ai_evaluation:
            label_text('▎AI 评语', 9)
            pdf.set_fill_color(*ACCENT_BG)
            y2 = pdf.get_y()
            pdf.set_font('cjk', '', 9)
            pdf.set_text_color(*DARK)
            pdf.multi_cell(page_w, 5.5, q.ai_evaluation)
            pdf.set_y(pdf.get_y() + 1)

        # Reference answer
        label_text('▎参考答案', 9)
        pdf.set_fill_color(*GREEN_BG)
        y3 = pdf.get_y()
        pdf.set_font('cjk', '', 9)
        pdf.set_text_color(*DARK)
        pdf.multi_cell(page_w, 5.5, q.reference_answer or '暂无')
        pdf.set_y(pdf.get_y() + 1)

        # Improvement
        if q.improvement_suggestion:
            label_text('▎改进建议', 9)
            pdf.set_fill_color(*ORANGE_BG)
            y4 = pdf.get_y()
            pdf.set_font('cjk', '', 9)
            pdf.set_text_color(*DARK)
            pdf.multi_cell(page_w, 5.5, q.improvement_suggestion)
            pdf.set_y(pdf.get_y() + 1)

        if idx < len(sorted_qs) - 1:
            divider()

    # ====== Resume Suggestions ======
    if interview.resume_suggestions:
        section_title('简历优化建议')
        body_text(interview.resume_suggestions)

    # ====== Footer ======
    pdf.ln(8)
    divider()
    pdf.set_font('cjk', '', 8)
    pdf.set_text_color(*GRAY)
    pdf.cell(page_w / 2, 5, '由 AI 模拟面试系统生成', align='L')
    pdf.cell(page_w / 2, 5, f'第 {{nb}} 页', align='R')
    pdf.set_text_color(*DARK)

    return pdf.output()


def _build_docx(interview: Interview, questions: list[InterviewQuestion]) -> bytes:
    """生成 Word (.docx) 格式报告 — python-docx，纯 Python 不依赖外部工具"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    doc = Document()
    scores = interview.dimension_scores or {}

    # 标题
    title = doc.add_heading('模拟面试报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 概览
    doc.add_heading('面试概览', 1)
    table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
    for row, (k, v) in zip(table.rows, [
        ('总体评分', f"{interview.total_score or '-'} 分 / 100"),
        ('难度级别', interview.difficulty),
        ('面试时间', interview.started_at.isoformat() if interview.started_at else 'N/A'),
        ('内容完整性', scores.get('content_completeness', '-')),
        ('专业度', scores.get('professionalism', '-')),
    ]):
        row.cells[0].text = k
        row.cells[1].text = str(v)

    # 能力分析
    if interview.ai_overview:
        doc.add_heading('综合评价', 1)
        doc.add_paragraph(interview.ai_overview)

    # 逐题详情
    doc.add_heading('逐题详情', 1)
    for q in sorted(questions, key=lambda x: x.order_index):
        sd = q.score_detail or {}
        doc.add_heading(f"第{q.order_index}题：{q.question_text[:50]}", 2)
        doc.add_paragraph(f"你的回答：{q.user_answer_transcript or '（未作答）'}")
        if q.ai_score is not None:
            doc.add_paragraph(f"评分：{q.ai_score} 分")
        if q.ai_evaluation:
            doc.add_paragraph(f"评语：{q.ai_evaluation}")
        doc.add_paragraph(f"参考答案：{q.reference_answer or '暂无'}")
        if q.improvement_suggestion:
            doc.add_paragraph(f"改进建议：{q.improvement_suggestion}")

    # 简历建议
    if interview.resume_suggestions:
        doc.add_heading('简历优化建议', 1)
        doc.add_paragraph(interview.resume_suggestions)

    doc.add_paragraph('')
    doc.add_paragraph('由 AI 模拟面试系统生成').alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_context(interview: Interview, questions: list[InterviewQuestion]) -> dict:
    """构建模板上下文"""
    scores = interview.dimension_scores or {}
    overview_data = {}

    # Parse AI overview for learning plan
    try:
        import json
        # ai_overview may contain a JSON with strengths/weaknesses/learning_plan
        if interview.ai_overview:
            # Try to extract structured data
            overview_data = {"gap_analysis": interview.ai_overview}
    except Exception:
        overview_data = {"gap_analysis": interview.ai_overview or "暂无分析"}

    q_list = []
    for q in sorted(questions, key=lambda x: x.order_index):
        score_detail = q.score_detail or {}
        # 格式化时长
        def _fmt_time(s: int | None) -> str:
            if s is None: return "-"
            m, sec = divmod(s, 60)
            return f"{m}:{sec:02d}" if m > 0 else f"{sec}秒"

        q_list.append({
            "index": q.order_index,
            "question_type_label": QUESTION_TYPE_LABELS.get(q.question_type, q.question_type),
            "question_text": q.question_text,
            "answer": q.user_answer_transcript or "（未回答）",
            "thinking_time": _fmt_time(q.thinking_duration_seconds),
            "answer_time": _fmt_time(q.duration_seconds),
            "content_score": score_detail.get("content_completeness", "-"),
            "professional_score": score_detail.get("professionalism", "-"),
            "expression_score": score_detail.get("expression", "-"),
            "star_score": score_detail.get("star_method", "-"),
            "total_score": q.ai_score or "-",
            "evaluation": q.ai_evaluation or "暂无评语",
            "reference_answer": q.reference_answer or "暂无参考答案",
            "improvement": q.improvement_suggestion or "暂无建议",
        })

    return {
        "position": "面试岗位",
        "interview_time": interview.started_at.isoformat() if interview.started_at else "N/A",
        "total_score": interview.total_score or "-",
        "difficulty": interview.difficulty,
        "content_score": scores.get("content_completeness", "-"),
        "professional_score": scores.get("professionalism", "-"),
        "expression_score": scores.get("expression", "-"),
        "star_score": scores.get("star_method", "-"),
        "gap_analysis": overview_data.get("gap_analysis", ""),
        "questions": q_list,
        "short_term": overview_data.get("short_term", []),
        "medium_term": overview_data.get("medium_term", []),
        "long_term": overview_data.get("long_term", []),
        "resume_suggestions": interview.resume_suggestions or "暂无建议",
        "generated_at": datetime.now().isoformat(),
    }


async def generate_document(
    db: AsyncSession,
    interview_id: uuid.UUID,
    doc_format: str,
    storage_dir: str,
) -> str:
    """生成文档并保存到本地文件"""
    result = await db.execute(select(Interview).where(Interview.id == interview_id))
    interview = result.scalar_one_or_none()
    if not interview:
        raise ValueError("Interview not found")

    q_result = await db.execute(
        select(InterviewQuestion).where(InterviewQuestion.interview_id == interview_id)
        .order_by(InterviewQuestion.order_index)
    )
    questions = q_result.scalars().all()

    # 获取简历和JD信息用于生成文件名
    candidate_name = ""
    job_position = ""
    try:
        r_result = await db.execute(select(Resume).where(Resume.id == interview.resume_id))
        resume = r_result.scalar_one_or_none()
        if resume and resume.parsed_data:
            candidate_name = (resume.parsed_data.get("basic") or {}).get("name", "") or ""
        # Fallback: use original filename (without extension) as candidate name
        if not candidate_name and resume:
            fname = resume.original_filename or ""
            candidate_name = os.path.splitext(fname)[0] if '.' in fname else fname
        j_result = await db.execute(select(JobDescription).where(JobDescription.id == interview.jd_id))
        jd = j_result.scalar_one_or_none()
        if jd and jd.parsed_data:
            job_position = (jd.parsed_data.get("position") or "").strip()
        # Fallback: use first line of JD raw text (truncated)
        if not job_position and jd:
            raw = (jd.raw_text or "").strip()
            job_position = raw.split('\n')[0][:30] if raw else ""
    except Exception:
        pass

    # Build professional filename: 模拟面试_姓名_岗位_日期.ext
    import re
    date_str = datetime.now().strftime("%Y%m%d")
    name_part = (candidate_name or "候选人").strip()
    position_part = (job_position or "面试岗位").strip()
    # 清洗：去路径特殊字符、UUID 前缀、多余空格
    name_part = re.sub(r'[\\/:*?"<>|]', '', name_part)
    name_part = re.sub(r'[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}[_-]*', '', name_part, flags=re.I)
    name_part = re.sub(r'\s+', '', name_part).strip('_') or "候选人"
    position_part = re.sub(r'[\\/:*?"<>|,]', '', position_part)
    position_part = re.sub(r'\s+', '', position_part).strip()[:20] or "面试"
    ext = doc_format
    filename = f"模拟面试_{name_part}_{position_part}_{date_str}.{ext}"

    if doc_format == "md":
        content = _build_markdown(interview, questions)
    elif doc_format == "html":
        content = _build_html(interview, questions)
    elif doc_format == "pdf":
        content = _build_pdf(interview, questions)
    elif doc_format == "docx":
        content = _build_docx(interview, questions)
    else:
        raise ValueError(f"Unsupported format: {doc_format}")

    # 保存文件
    user_dir = Path(storage_dir) / str(interview.user_id) / "reports"
    user_dir.mkdir(parents=True, exist_ok=True)
    filepath = user_dir / filename

    binary_formats = ("pdf", "docx")
    mode = "wb" if doc_format in binary_formats else "w"
    encoding = None if doc_format in binary_formats else "utf-8"
    with open(filepath, mode, encoding=encoding) as f:
        f.write(content)

    # 创建文档记录
    doc_record = InterviewDocument(
        interview_id=interview_id,
        format=doc_format,
        file_path=str(filepath),
        file_size=os.path.getsize(filepath),
    )
    db.add(doc_record)
    await db.commit()

    return str(filepath)
